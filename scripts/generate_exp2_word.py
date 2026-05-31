# -*- coding: utf-8 -*-
"""Generate Experiment 2 (LR(1) syntax analysis) Word report on desktop."""
from __future__ import annotations

import sys
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

DESKTOP = Path(r"c:\Users\Cc\Desktop")
OUTPUT_DOCX = DESKTOP / "实验报告二.docx"
OUTPUT_DOC = DESKTOP / "实验报告二.doc"
IMG_CALL = DESKTOP / "_exp2_call_graph.png"
IMG_FLOW = DESKTOP / "_exp2_flow_chart.png"
PROJECT_ROOT = Path(__file__).resolve().parent.parent
JAVA_ROOT = PROJECT_ROOT / "src" / "main" / "java"
TEST_OUTPUT_DIR = PROJECT_ROOT / "docs" / "test-output"
SAMPLES_DIR = PROJECT_ROOT / "samples"

# 与 export-test-outputs.ps1 导出的 7 个样例一一对应
SAMPLE_CASES = [
    ("sample1_basic", "测试组 1：基础声明与赋值（sample1_basic）"),
    ("sample2_if_else", "测试组 2：if-else 与双字符运算符（sample2_if_else）"),
    ("sample3_loop", "测试组 3：while / do 关键字（sample3_loop）"),
    ("sample4_break_error", "测试组 4：break 关键字（sample4_break_error）"),
    ("sample5_array", "测试组 5：数组声明与下标（sample5_array）"),
    ("sample6_float", "测试组 6：实型常数（sample6_float）"),
    ("sample7_bool", "测试组 7：bool 与布尔常数（sample7_bool）"),
]


def java_lines(rel: str, start: int, end: int) -> str:
    """Read inclusive line range from project Java source (1-based)."""
    path = JAVA_ROOT / rel
    lines = path.read_text(encoding="utf-8").splitlines()
    return "\n".join(lines[start - 1 : end])


def read_sample_source(base_name: str) -> str:
    path = SAMPLES_DIR / f"{base_name}.src"
    return path.read_text(encoding="utf-8").rstrip()


def read_output_text(txt_path: Path) -> str:
    raw = txt_path.read_bytes()
    if raw.startswith(b"\xff\xfe") or raw.startswith(b"\xfe\xff"):
        return raw.decode("utf-16")
    if raw.startswith(b"\xef\xbb\xbf"):
        return raw.decode("utf-8-sig")
    return raw.decode("utf-8")


def parse_exp1_from_output(txt_path: Path) -> tuple[list[str], list[str]]:
    """从 docs/test-output/*.txt 解析「实验一」二元组与符号表（至实验二之前）。"""
    tokens: list[str] = []
    symbols: list[str] = []
    in_tokens = False
    in_symbols = False
    for line in read_output_text(txt_path).splitlines():
        if line.startswith("========== 实验二"):
            break
        if line.strip() == "--- (单词名称, 单词类别) ---":
            in_tokens, in_symbols = True, False
            continue
        if line.strip() == "--- 符号表（按出现顺序） ---":
            in_tokens, in_symbols = False, True
            continue
        if line.startswith("--- "):
            continue
        if in_tokens and line.startswith("("):
            tokens.append(line)
        if in_symbols and line.strip().startswith("("):
            symbols.append(line.strip())
    return tokens, symbols


def load_test_cases_from_exports() -> list[tuple[str, str, list[str], str, str]]:
    """(标题, 源程序, 二元组行, 符号表说明, 输出来源文件)"""
    cases: list[tuple[str, str, list[str], str, str]] = []
    for base, title in SAMPLE_CASES:
        out_file = TEST_OUTPUT_DIR / f"{base}.txt"
        if not out_file.exists():
            raise FileNotFoundError(f"缺少导出文件: {out_file}，请先运行 scripts/export-test-outputs.ps1")
        tokens, symbols = parse_exp1_from_output(out_file)
        source = read_sample_source(base)
        symbol_note = "符号表：\n" + "\n".join(symbols) if symbols else "符号表：(无标识符)"
        cases.append((title, source, tokens, symbol_note, f"docs/test-output/{base}.txt"))
    return cases


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name: str = "宋体", size: int = 12, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, "黑体", 16 if level == 1 else 14, bold=True)


def add_para(doc: Document, text: str, bold: bool = False, indent: bool = False) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    set_run_font(p.add_run(text), bold=bold)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        set_run_font(doc.add_paragraph(style="List Bullet").add_run(item))


def add_code_block(doc: Document, text: str) -> None:
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "D9E2F3")
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def add_picture(doc: Document, path: Path, caption: str, width_in: float = 6.2) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cap.add_run(caption), size=10)
    doc.add_paragraph()


def _setup_cjk_font() -> None:
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "SimSun", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def _draw_box(ax, x, y, w, h, text, fc="#E8F4FC", ec="#2E75B6", fs=9):
    box = FancyBboxPatch(
        (x - w / 2, y - h / 2),
        w,
        h,
        boxstyle="round,pad=0.03",
        linewidth=1.2,
        edgecolor=ec,
        facecolor=fc,
    )
    ax.add_patch(box)
    ax.text(x, y, text, ha="center", va="center", fontsize=fs, linespacing=1.25)


def _draw_diamond(ax, x, y, w, h, text, fc="#FFF4E5", ec="#C55A11"):
    diamond = plt.Polygon(
        [(x, y + h / 2), (x + w / 2, y), (x, y - h / 2), (x - w / 2, y)],
        closed=True,
        facecolor=fc,
        edgecolor=ec,
        linewidth=1.2,
    )
    ax.add_patch(diamond)
    ax.text(x, y, text, ha="center", va="center", fontsize=9)


def _arrow(ax, x1, y1, x2, y2, text: str | None = None):
    ax.add_patch(
        FancyArrowPatch(
            (x1, y1),
            (x2, y2),
            arrowstyle="->",
            mutation_scale=12,
            linewidth=1.0,
            color="#444444",
            connectionstyle="arc3,rad=0.0",
        )
    )
    if text:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.12, text, ha="center", va="bottom", fontsize=8)


def generate_call_graph(path: Path) -> None:
    """分层树形布局，避免左右两支节点重叠。"""
    _setup_cjk_font()
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")

    # 左：编译流水线；右：演示输出（分列）
    nodes = [
        (6.0, 7.2, 3.2, 0.55, "CompilerApplication.main"),
        (2.2, 5.8, 2.4, 0.5, "ExperimentDemo\n.loadSource"),
        (5.0, 5.8, 2.4, 0.5, "CompilerPipeline\n.compile"),
        (9.8, 5.8, 2.4, 0.5, "ExperimentDemo\n.printReport"),
        (5.0, 4.2, 2.2, 0.5, "LexerImpl.analyze"),
        (9.8, 4.2, 2.6, 0.45, "LexicalReporter\n.formatTokenTuples"),
        (9.8, 3.0, 2.6, 0.45, "LexicalReporter\n.formatSymbolTable"),
        (1.2, 2.0, 2.0, 0.45, "SourceBuffer"),
        (3.4, 2.0, 2.0, 0.45, "KeywordTable\n.isKeyword"),
        (5.6, 2.0, 2.4, 0.45, "ScopedSymbolTable\n.registerIdentifier"),
        (7.8, 2.0, 2.6, 0.45, "LexicalDeclarationBinder\n.bindDeclarations"),
        (7.8, 0.8, 2.4, 0.45, "ScopedSymbolTable\n.bindDeclaration"),
    ]
    for x, y, w, h, t in nodes:
        _draw_box(ax, x, y, w, h, t)

    _arrow(ax, 6.0, 6.9, 2.2, 6.05)
    _arrow(ax, 6.0, 6.9, 5.0, 6.05)
    _arrow(ax, 6.0, 6.9, 9.8, 6.05)
    _arrow(ax, 5.0, 5.55, 5.0, 4.45)
    _arrow(ax, 9.8, 5.55, 9.8, 4.65)
    _arrow(ax, 9.8, 5.55, 9.8, 3.35)
    _arrow(ax, 5.0, 3.95, 1.2, 2.25)
    _arrow(ax, 5.0, 3.95, 3.4, 2.25)
    _arrow(ax, 5.0, 3.95, 5.6, 2.25)
    _arrow(ax, 5.0, 3.95, 7.8, 2.25)
    _arrow(ax, 7.8, 1.75, 7.8, 1.05)

    ax.set_title("图1  词法分析模块函数调用关系图", fontsize=12, pad=10)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def generate_flow_chart(path: Path) -> None:
    """纵向主流程 + 独立分支说明框，避免多框同层重叠。"""
    _setup_cjk_font()
    fig, ax = plt.subplots(figsize=(7.5, 12))
    ax.set_xlim(0, 7.5)
    ax.set_ylim(0, 12)
    ax.axis("off")

    cx = 2.6
    bx = 5.8

    _draw_box(ax, cx, 11.0, 3.8, 0.55, "开始：读入 sourceCode")
    _draw_box(ax, cx, 9.8, 3.8, 0.7, "创建 SourceBuffer、符号表\ntokens / errors 列表")
    _draw_diamond(ax, cx, 8.5, 3.2, 0.9, "index < length ?")
    _draw_box(ax, cx, 6.8, 3.8, 0.55, "追加 <EOF> token")
    _draw_box(ax, cx, 5.6, 3.8, 0.65, "LexicalDeclarationBinder\n.bindDeclarations")
    _draw_box(ax, cx, 4.4, 3.8, 0.55, "返回 LexicalAnalyzerResult")

    branch_text = (
        "循环体内分支：\n"
        "· 空白 → 更新行列，continue\n"
        "· // 、/* */ → 跳过注释\n"
        "  （未闭合 → LexicalError）\n"
        "· 字母/_ → 标识符/关键字/布尔\n"
        "· 数字 → 整数或实数\n"
        "· 双/单字符运算符\n"
        "· 分隔符 { } ( ) [ ] ;\n"
        "· 其它 → Illegal character\n"
        "  记录错误后跳过 1 字符"
    )
    _draw_box(ax, bx, 8.2, 3.2, 2.6, branch_text, fc="#F5F5F5", ec="#666666", fs=8)

    _arrow(ax, cx, 10.75, cx, 10.15)
    _arrow(ax, cx, 9.45, cx, 8.95)
    _arrow(ax, cx + 1.6, 8.5, bx - 1.6, 8.5)
    _arrow(ax, bx, 6.9, cx - 1.5, 8.95)
    _arrow(ax, cx, 8.05, cx, 7.1)
    ax.text(cx + 0.28, 7.55, "否", fontsize=8)
    ax.text(bx - 0.5, 7.5, "是", fontsize=8)
    _arrow(ax, cx, 6.55, cx, 5.95)
    _arrow(ax, cx, 5.25, cx, 4.65)

    ax.set_title("图2  LexerImpl.analyze 总体执行流程图", fontsize=12, pad=10)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


API_ROWS = [
    ["Lexer", "LexicalAnalyzerResult analyze(String sourceCode)",
     "sourceCode：待扫描的完整源程序文本",
     "LexicalAnalyzerResult", "词法分析对外接口"],
    ["LexerImpl", "LexicalAnalyzerResult analyze(String sourceCode)",
     "sourceCode：源程序；内部维护 index/line/column",
     "tokens、errors、symbolsInOrder", "主扫描：空白/注释/单词/错误"],
    ["SourceBuffer", "char charAt(int index)\nString substring(int start, int end)",
     "index/start/end：在 char[] 上的下标",
     "char 或子串", "缓冲区按索引读字符，避免重复 substring"],
    ["KeywordTable", "boolean isKeyword(String lexeme)",
     "lexeme：标识符拼写",
     "true 表示保留字", "区分 if/int 与普通标识符"],
    ["ScopedSymbolTable", "void registerIdentifier(String name)",
     "name：标识符名",
     "void", "首次出现时按顺序加入符号表，类型暂为「待声明」"],
    ["ScopedSymbolTable", "void bindDeclaration(String name, String typeName, SymbolKind kind, Integer arrayLength)",
     "name/typeName/kind/arrayLength：声明得到的类型信息",
     "void", "将已登记标识符更新为 int/float/bool/int[] 等"],
    ["LexicalDeclarationBinder", "void bindDeclarations(List<Token> tokens, ScopedSymbolTable symbolTable)",
     "tokens：扫描结果；symbolTable：待填类型的符号表",
     "void", "匹配 type id; 与 type[num] id; 后调用 bindDeclaration"],
    ["LexicalReporter", "List<String> formatTokenTuples(List<Token> tokens)",
     "tokens：含 EOF 的 token 列表",
     "List<String>", "生成 (单词, 类别) 二元组行，过滤 EOF"],
    ["LexicalReporter", "List<String> formatSymbolTable(List<Symbol> symbols)",
     "symbols：symbolsInOrder()",
     "List<String>", "生成 (名, 类型, 种类) 符号表行"],
    ["ExperimentDemo", "void printReport(PipelineReport report, LR1Parser parser, String[] args)",
     "report：含 lexicalResult；用于控制台实验一输出",
     "void", "打印二元组、符号表、词法错误"],
]


def build_document() -> Document:
    generate_call_graph(IMG_CALL)
    generate_flow_chart(IMG_FLOW)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.5)

    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(t.add_run("编译系统设计实践实验报告"), "黑体", 22, True)
    doc.add_paragraph()
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(t2.add_run("实验二：LR(1) 语法分析程序实验"), "黑体", 18, True)
    doc.add_paragraph()
    for label in ["学    号：", "姓    名：", "年    级：", "学    院：", "专    业："]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(f"{label}  ________________"), size=14)
    add_table(doc, ["小组成员", "学号", "分工"], [["", "", "LR(1) 语法分析、实验报告"], ["", "", ""]])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("实验时间：2025—2026 学年第二学期"), size=12)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p2.add_run("授 课 教 师：________________"), size=12)
    doc.add_page_break()

    add_heading(doc, "一、实验目的与任务", 1)
    add_para(
        doc,
        "从源程序识别关键字、标识符、常数、运算符、分隔符；以 (单词名称, 单词类别) 输出；"
        "删除注释与空白；定位词法错误；建立并按顺序打印符号表。本实验实现位于包 edu.groupname.compiler.lexer。",
        indent=True,
    )

    add_heading(doc, "二、功能描述", 1)
    add_para(
        doc,
        "本程序（LexerImpl 及配套类）实现课程实验一词法分析器，具体功能如下：",
        indent=True,
    )
    add_bullets(
        doc,
        [
            "分隔单词并以二元组 (单词名称, 单词类别) 输出，类别由 TokenCategoryLabel 转为中文（关键字、标识符、整数、实数、布尔常数、运算符、分隔符）。",
            "删除 // 行注释、/* */ 块注释及空格、制表符、换行，不生成对应 token。",
            "建立 ScopedSymbolTable：标识符首次出现时 registerIdentifier；扫描结束后 LexicalDeclarationBinder 根据 int/float/bool 声明绑定类型，并按顺序打印。",
            "词法错误：非法字符、未闭合块注释、实数点后无数字等，通过 LexicalError 携带 SourcePosition（行、列）。",
            "演示入口 ExperimentDemo 在 CompilerApplication 运行时输出「实验二：LR(1) 语法分析」小节。",
        ],
    )

    add_heading(doc, "三、程序结构描述", 1)
    add_heading(doc, "3.1 函数/方法说明（参数含义、返回值、功能）", 2)
    add_table(
        doc,
        ["类", "函数原型", "参数含义", "返回值", "功能"],
        API_ROWS,
        col_widths=[2.2, 3.2, 3.5, 2.2, 2.8],
    )

    add_heading(doc, "3.2 函数调用关系图", 2)
    add_para(doc, "下图与工程调用链一致：main 经 CompilerPipeline 调用 LexerImpl.analyze，再经 ExperimentDemo 调用 LexicalReporter 输出。", indent=True)
    add_picture(doc, IMG_CALL, "图1  词法分析模块函数调用关系图")

    add_heading(doc, "3.3 程序总体执行流程图", 2)
    add_para(doc, "下图对应 LexerImpl.analyze 中 while 循环及各分支（与源码第 28～168 行一致）。", indent=True)
    add_picture(doc, IMG_FLOW, "图2  LexerImpl.analyze 总体执行流程图")

    add_heading(doc, "四、符号表的设计和结构", 1)
    add_heading(doc, "4.1 设计思想", 2)
    add_bullets(
        doc,
        [
            "registerIdentifier(String name)：在 LexerImpl 识别到普通标识符时调用，仅首次出现加入 definitionOrder。",
            "初始 Symbol 的 typeName 为「待声明」，kind 为 VARIABLE。",
            "bindDeclarations 遍历 token，识别 KEYWORD 为 int/float/bool 后的 IDENTIFIER 或 int[num] id 模式，调用 bindDeclaration 更新类型（含 int[] 与 arrayLength）。",
        ],
    )
    add_heading(doc, "4.2 符号表项结构（Symbol record）", 2)
    add_table(
        doc,
        ["字段", "含义", "示例"],
        [
            ["name", "标识符名", "a"],
            ["typeName", "int / float / bool / int[] / 待声明", "int"],
            ["kind", "VARIABLE 或 ARRAY", "VARIABLE"],
            ["arrayLength", "数组声明中的长度", "10"],
        ],
    )
    add_para(doc, "输出示例：(a, int, VARIABLE)；(arr, int[], ARRAY, len=10)。", indent=True)

    add_heading(doc, "五、有限状态自动机与缓冲区（简要）", 1)
    add_para(doc, "SourceBuffer 提供字符缓冲区；LexerImpl 用直接编码实现 DFA（START、注释态、IN_ID、IN_NUM/IN_REAL、ERROR 等），与教材第 3 章一致。", indent=True)

    add_heading(doc, "六、程序关键源代码", 1)
    add_para(doc, "以下摘录自本工程 src/main/java/edu/groupname/compiler/lexer 与 symbol 包，与仓库代码一致。", indent=True)

    lexer = "edu/groupname/compiler/lexer/LexerImpl.java"
    add_heading(doc, "6.1 LexerImpl.analyze 入口（第 17～29 行）", 2)
    add_code_block(doc, java_lines(lexer, 17, 29))

    add_heading(doc, "6.2 空白与注释处理（第 31～79 行）", 2)
    add_code_block(doc, java_lines(lexer, 31, 79))

    add_heading(doc, "6.3 标识符与关键字（第 84～99 行）", 2)
    add_code_block(doc, java_lines(lexer, 84, 99))

    add_heading(doc, "6.4 常数、运算符、分隔符与错误（第 102～155 行）", 2)
    add_code_block(doc, java_lines(lexer, 102, 155))

    add_heading(doc, "6.5 结束与返回（第 157～168 行）", 2)
    add_code_block(doc, java_lines(lexer, 157, 168))

    add_heading(doc, "6.6 SourceBuffer.java（第 6～26 行）", 2)
    add_code_block(doc, java_lines("edu/groupname/compiler/lexer/SourceBuffer.java", 6, 26))

    add_heading(doc, "6.7 LexicalDeclarationBinder.bindDeclarations（第 20～50 行）", 2)
    add_code_block(doc, java_lines("edu/groupname/compiler/lexer/LexicalDeclarationBinder.java", 20, 50))

    add_heading(doc, "6.8 LexicalReporter.formatTuple（第 15～24 行）", 2)
    add_code_block(doc, java_lines("edu/groupname/compiler/lexer/LexicalReporter.java", 15, 24))

    add_heading(doc, "6.9 ScopedSymbolTable.registerIdentifier / bindDeclaration（第 30～55 行）", 2)
    add_code_block(doc, java_lines("edu/groupname/compiler/symbol/ScopedSymbolTable.java", 30, 55))

    add_heading(doc, "七、测试数据与运行结果（≥3 组）", 1)
    add_para(
        doc,
        "以下 7 组数据均来自工程样例源程序 samples/*.src，经 CompilerApplication 运行后由 "
        "scripts/export-test-outputs.ps1 导出；词法部分摘自 docs/test-output/*.txt 中「实验二：LR(1) 语法分析」"
        "章节（与控制台输出一致），未另行编造。",
        indent=True,
    )
    add_para(
        doc,
        "复现命令（Compiler 目录）：powershell -ExecutionPolicy Bypass -File .\\scripts\\export-test-outputs.ps1",
        indent=True,
    )
    add_para(
        doc,
        "注：右括号「)」的二元组格式为 ( , 分隔符) 时写作「(), 分隔符)」（左括号与 lexeme 相连），与 LexicalReporter.formatTuple 输出一致。",
        indent=True,
    )

    for title, source, token_lines, symbol_note, out_ref in load_test_cases_from_exports():
        base = out_ref.split("/")[-1].replace(".txt", "")
        add_heading(doc, title, 2)
        add_para(doc, f"输入（samples/{base}.src）：", bold=True)
        add_code_block(doc, source)
        add_para(doc, f"运行结果（摘自 {out_ref}「实验一」章节）：", bold=True)
        add_code_block(doc, "\n".join(token_lines))
        add_para(doc, symbol_note, indent=True)

    add_heading(doc, "八、实验总结", 1)
    add_bullets(
        doc,
        [
            "完成了任务书要求的二元组输出、注释/空白删除、错误定位、符号表建立与顺序打印，实现类为 LexerImpl、SourceBuffer、LexicalDeclarationBinder、ScopedSymbolTable、LexicalReporter。",
            "采用字符缓冲区与 DFA 直接编码，与教材词法分析原理一致；双字符运算符、布尔字面量 true/false 与关键字分流与源码一致。",
            "测试数据见 docs/test-output 下 7 个样例导出文件之「实验一」章节；LexerImplTest 单元测试全部通过。",
            "词法符号表为后续 CompilerPipeline 中 LR(1) 语法分析提供 token 序列与标识符类型信息。",
        ],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("教材：编译原理（第 2 版），Alfred V. Aho 等著，赵建华等译，机械工业出版社"), size=10)
    return doc


def save_as_doc(docx_path: Path, doc_path: Path) -> bool:
    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(doc_path.resolve()), FileFormat=0)
        doc.Close(False)
        word.Quit()
        return True
    except Exception as exc:
        print(f"Word COM failed: {exc}", file=sys.stderr)
        return False


def main() -> int:
    _setup_cjk_font()
    DESKTOP.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(OUTPUT_DOCX))
    print(f"Saved: {OUTPUT_DOCX}")

    ok = save_as_doc(OUTPUT_DOCX, OUTPUT_DOC)
    for p in (IMG_CALL, IMG_FLOW):
        try:
            p.unlink(missing_ok=True)
        except OSError:
            pass
    if ok:
        try:
            TEMP_DOCX.unlink(missing_ok=True)
        except OSError:
            pass
        print(f"Updated: {OUTPUT_DOC}")
        return 0

    fallback_docx = DESKTOP / "实验报告二.docx"
    doc.save(str(fallback_docx))
    print(f"WPS/Word 未转换 .doc，已保存: {fallback_docx}", file=sys.stderr)
    print("请关闭已打开的 实验报告一_更新.doc 后重新运行本脚本，或直接用 .docx 打开。", file=sys.stderr)
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
