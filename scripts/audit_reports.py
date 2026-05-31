# -*- coding: utf-8 -*-
"""Audit experiment reports vs code, test-output, and task book."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

DESK = Path(r"c:\Users\Cc\Desktop")
JAVA = Path(r"d:\fzu\PrinciplesOfCompilation\Compiler\src\main\java")
TEST = Path(r"d:\fzu\PrinciplesOfCompilation\Compiler\docs\test-output")
SAMPLES = Path(r"d:\fzu\PrinciplesOfCompilation\Compiler\samples")
TMP = Path(__file__).resolve().parent / "_audit"


def open_doc(path: Path) -> Document:
    if path.suffix.lower() == ".docx":
        return Document(str(path))
    import win32com.client

    TMP.mkdir(exist_ok=True)
    out = TMP / f"{path.stem}.docx"
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    doc = word.Documents.Open(str(path.resolve()))
    doc.SaveAs2(str(out.resolve()), FileFormat=16)
    doc.Close(False)
    word.Quit()
    return Document(str(out))


def mono_text(doc: Document) -> str:
    parts: list[str] = []
    for para in doc.paragraphs:
        for run in para.runs:
            fn = run.font.name or ""
            if ("Courier" in fn or "Consolas" in fn) and run.text:
                parts.append(run.text)
    return "".join(parts)


def body_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def find_report(num: str) -> Path | None:
    base = f"\u5b9e\u9a8c\u62a5\u544a{num}"
    for ext in (".doc", ".docx"):
        p = DESK / f"{base}{ext}"
        if p.exists():
            return p
    p = DESK / "\u65b0\u5efa\u6587\u4ef6\u5939" / f"{base}.docx"
    return p if p.exists() else None


def java_slice(rel: str, start: int, end: int) -> str:
    lines = (JAVA / rel).read_text(encoding="utf-8").splitlines()
    return "\n".join(lines[start - 1 : end])


def verify_excerpt(code: str, rel: str, start: int, end: int, label: str) -> dict:
    snippet = java_slice(rel, start, end)
    sig = [ln.strip() for ln in snippet.splitlines() if ln.strip() and not ln.strip().startswith("//")][:8]
    hits = [ln for ln in sig if ln in code]
    ratio = len(hits) / max(len(sig), 1)
    return {"label": label, "hits": len(hits), "total": len(sig), "ratio": ratio, "ok": ratio >= 0.5}


def read_test_section(path: Path, marker: str, stop: str | None = None) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    i = text.find(marker)
    if i < 0:
        return ""
    chunk = text[i:]
    if stop:
        j = chunk.find(stop)
        if j > 0:
            chunk = chunk[:j]
    return chunk


def main() -> int:
    reports = {"\u4e00": find_report("\u4e00"), "\u4e8c": find_report("\u4e8c"), "\u4e09": find_report("\u4e09")}
    print("REPORTS:")
    for k, p in reports.items():
        print(f"  {k}: {p}")

    excerpts = {
        "\u4e00": [
            ("LexerImpl.analyze", "edu/groupname/compiler/lexer/LexerImpl.java", 17, 80),
            ("SourceBuffer", "edu/groupname/compiler/lexer/SourceBuffer.java", 1, 27),
            ("LexicalDeclarationBinder", "edu/groupname/compiler/lexer/LexicalDeclarationBinder.java", 20, 51),
            ("LexicalReporter", "edu/groupname/compiler/lexer/LexicalReporter.java", 15, 37),
        ],
        "\u4e8c": [
            ("LR1Parser.parse", "edu/groupname/compiler/parser/LR1Parser.java", 46, 146),
            ("buildParsingTable", "edu/groupname/compiler/parser/LR1Parser.java", 196, 251),
            ("closure/goTo", "edu/groupname/compiler/parser/LR1Parser.java", 254, 307),
            ("ParsingTable", "edu/groupname/compiler/parser/lr1/ParsingTable.java", 14, 61),
            ("ParseReporter", "edu/groupname/compiler/parser/ParseReporter.java", 14, 28),
        ],
        "\u4e09": [
            ("SemanticReduceHandler", "edu/groupname/compiler/parser/SemanticReduceHandler.java", 55, 120),
            ("LR1Parser reduce", "edu/groupname/compiler/parser/LR1Parser.java", 95, 135),
            ("IrBuilder", "edu/groupname/compiler/ir/IrBuilder.java", 1, 57),
            ("IrReporter", "edu/groupname/compiler/ir/IrReporter.java", 1, 39),
            ("SemanticAnalyzerImpl", "edu/groupname/compiler/semantic/SemanticAnalyzerImpl.java", 25, 80),
            ("IRGeneratorImpl", "edu/groupname/compiler/ir/IRGeneratorImpl.java", 1, 16),
            ("ScopedSymbolTable", "edu/groupname/compiler/symbol/ScopedSymbolTable.java", 1, 55),
        ],
    }

    test_expect = {
        "\u4e00": [
            ("sample1_basic", "\u5b9e\u9a8c\u4e00", "\u5b9e\u9a8c\u4e8c", ["(int,", "(a,", "VARIABLE"]),
            ("sample2_if_else", "\u5b9e\u9a8c\u4e00", "\u5b9e\u9a8c\u4e8c", ["(if,", ">="]),
            ("sample3_loop", "\u5b9e\u9a8c\u4e00", "\u5b9e\u9a8c\u4e8c", ["(while,", "(do,"]),
        ],
        "\u4e8c": [
            ("sample1_basic", "\u5b9e\u9a8c\u4e8c", "\u5b9e\u9a8c\u4e09", ["accepted = true", "ACTION(", "\u72b6\u6001\u6808="]),
            ("sample2_if_else", "\u5b9e\u9a8c\u4e8c", "\u5b9e\u9a8c\u4e09", ["accepted = true", "acc"]),
            ("sample3_loop", "\u5b9e\u9a8c\u4e8c", "\u5b9e\u9a8c\u4e09", ["accepted = true", "while"]),
        ],
        "\u4e09": [
            ("sample2_if_else", "\u5b9e\u9a8c\u4e09", None, ["(GE,", "IF_FALSE_GOTO", "\u72b6\u6001\u6808="]),
            ("sample3_loop", "\u5b9e\u9a8c\u4e09", None, ["(LABEL,", "IF_FALSE_GOTO"]),
            ("sample5_array", "\u5b9e\u9a8c\u4e09", None, ["ARRAY_LOAD", "ARRAY_STORE"]),
        ],
    }

    for key in ("\u4e00", "\u4e8c", "\u4e09"):
        p = reports[key]
        if not p:
            print(f"\nMISSING report {key}")
            continue
        doc = open_doc(p)
        full = body_text(doc)
        code = mono_text(doc)
        print(f"\n{'='*60}\nREPORT {key}: {p.name} size={p.stat().st_size}\n{'='*60}")
        print("CODE blocks chars:", len(code))
        print("KmՋ�~ count:", full.count("\u6d4b\u8bd5\u7ec4"))
        # excerpts
        print("--- Source excerpt verification ---")
        for item in excerpts[key]:
            r = verify_excerpt(code, item[1], item[2], item[3], item[0])
            status = "OK" if r["ok"] else "FAIL"
            print(f"  [{status}] {r['label']}: {r['hits']}/{r['total']} signature lines")
        # test output cross-check (report contains key strings from test-output)
        print("--- Test output cross-check (report vs test-output) ---")
        for sample, start, stop, needles in test_expect[key]:
            tout = read_test_section(TEST / f"{sample}.txt", start, stop)
            rep_sample = sample in full
            hits = sum(1 for n in needles if n in full)
            tout_hits = sum(1 for n in needles if n in tout)
            print(f"  {sample}: in_report={rep_sample} needles_in_report={hits}/{len(needles)} in_testout={tout_hits}/{len(needles)}")
        # contamination
        if key == "\u4e00":
            print("  contamination LexerImpl-only:", full.count("LR1Parser"), "LR1Parser mentions")
        if key == "\u4e8c":
            print("  contamination:", "LexerImpl", full.count("LexerImpl"), "ACTION", full.count("ACTION"))
        if key == "\u4e09":
            print("  contamination:", "͋�l�NCQ�~;NSO", full.count("\u4e8c\u5143\u7ec4") > 5)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
